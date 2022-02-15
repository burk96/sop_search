from pathlib import Path
import datetime as dt
import regex as re
from docx import Document
import pandas as pd

# SOP specific parameters
# Match files with a 4 digit code and possibly a character and/or space before a hyphen
SOP_REGEX = r"^(\d{4}\w?\s?)[-](.)+"
SOP_PATH = r"P:\SCRIPTS\ALL Files"


def sop_search(search_regex, search_path):
    """Search and parses SOP folder"""
    sops = []

    for path in Path(search_path).rglob("*.docx"):
        if (
            re.search(search_regex, path.name)
            and "Archive" not in path.parts
            and Path.is_file(path)
        ):
            # After number in foldername
            department = re.findall(r"((?<=\d{4}\w?\s?[-]\s?).*)", path.parts[3])[
                0
            ].strip()
            # 4 digits at start of string, 0 or 1 chars
            number = re.findall(r"^\d{4}\w?", path.name)[0]
            # Match everything after number
            title = re.findall(r"((?<=\d{4}\w?\s?[-]\s?).*)", path.stem)[0].strip()
            # Create an excel useable hyperlink for path
            hyperlink = f'=HYPERLINK("{path}", "{path}")'
            last_revision_date = last_revision_date_from_docx(path)

            row = {
                "Department": department,
                "Number": number,
                "File Name/Title": title,
                "Link to documents": hyperlink,
                "Last Revision Date": last_revision_date.strftime("%Y-%m-%d")
                if last_revision_date
                else None,
            }

            sops.append(row)
    return sops


def date_finder(file_obj):
    """Searches a file obj for date"""
    # Matches dates like M/D/YY, MM/DD/YYYY, and some permutations
    date_re = r"((0?[1-9]|1[0-2])[\/](0?[1-9]|[12]\d|3[01])[\/]((19|20)\d{2}|\d{2}))"
    dates = []

    # I know, this nesting is a bit much right? 😳
    for table in file_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if re.search(date_re, para.text):
                        date = re.findall(date_re, para.text)[0][0]
                        # Format YY to YYYY and convert to datetime
                        if len(date.split("/")[2]) == 2:
                            date_split = date.split("/")
                            date_split[2] = "20" + date_split[2]
                            date = "/".join(date_split)
                        date = dt.datetime.strptime(date, "%m/%d/%Y")
                        dates.append(date)
    return dates


def last_revision_date_from_docx(docx_path):
    """Searches tables in a .docx file for the last_revision date"""
    found_dates = None
    try:
        # Last revision date can be in document body or footer
        document = Document(docx_path)
        section = document.sections[0]
        footer = section.footer
        found_dates = [*date_finder(document), *date_finder(footer)]

    except IOError:
        print("Could not open file")

    # Given the somewhat unpredictableness of where the last revision date is stored
    # This script assumes the newest date in a table is the last revision date
    # This appears to be a safe assumption given all the files I've checked
    # But there is no gurantee
    return max(found_dates) if found_dates else None


def export_to_excel(files_dict, filename, sheetname):
    """Exports a list of dictionaries to excel"""
    d_f = pd.DataFrame(files_dict)

    # Autoformat column width
    # Modified from StackOverflow https://stackoverflow.com/a/40535454
    # Also pylint gives a false error on the line below, and black keeps expanding it 🤦‍♀️
    # fmt: off
    writer = pd.ExcelWriter(filename, engine="xlsxwriter")  # pylint: disable=abstract-class-instantiated
    # fmt: on
    d_f.to_excel(writer, sheet_name=sheetname, index=False)
    worksheet = writer.sheets[sheetname]
    for idx, col in enumerate(d_f):
        series = d_f[col]
        # len of largest item vs len of col name
        max_len = (
            max(
                (
                    series.astype(str).map(len).max(),
                    len(str(series.name)),
                )
            )
            + 1
        )
        # TODO: *Properly* fix Link to Docs coming out too large
        max_len = min(max_len, 111)
        worksheet.set_column(idx, idx, max_len)
    writer.save()


export_to_excel(sop_search(SOP_REGEX, SOP_PATH), "SOPS.xlsx", "SOPs")
